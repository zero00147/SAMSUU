"""Document ingest: extract text, split into sections by heading.

Sections are the unit of attachment. A 12.5k-token PRD cannot fit in an 8k window,
but §7.1 on its own is ~400 tokens — so the user attaches the section the current
task needs rather than the whole file.

Splitting is deterministic (heading regexes, no embeddings). On a numbered spec that
beats semantic retrieval and costs no extra RAM, which matters on an 8 GB machine.
"""

import io
import re

# Markdown headings: ## 7.1 Proxy Bidding
_MD_HEADING = re.compile(r"^(#{1,4})\s+(.+?)\s*$")

# Numbered headings in extracted PDF/DOCX text: "3.1 Item Listing Engine",
# "5. Technical Architecture". Requires a title after the number so that list items
# ("1. Media Assets: minimum 3 images…") and prose ("cut 3. 5 inches") don't match.
_NUM_HEADING = re.compile(r"^\s*(\d+(?:\.\d+){0,2})\.?\s+([A-Z][^\n]{2,80})$")

MAX_CHARS = 400_000          # ~130k tokens; refuse anything larger
FALLBACK_CHUNK_CHARS = 4_000  # used when a document has no detectable headings


class ExtractError(Exception):
    pass


def extract_text(filename: str, data: bytes) -> str:
    """Plain text from md/txt/pdf/docx."""
    name = filename.lower()

    if name.endswith((".md", ".txt", ".markdown", ".rst")):
        try:
            return data.decode("utf-8")
        except UnicodeDecodeError:
            return data.decode("latin-1", errors="replace")

    if name.endswith(".pdf"):
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ExtractError("pypdf is not installed")
        try:
            reader = PdfReader(io.BytesIO(data))
            pages = [p.extract_text() or "" for p in reader.pages]
        except Exception as exc:
            raise ExtractError(f"could not read PDF: {exc}")
        text = "\n\n".join(pages)
        if not text.strip():
            raise ExtractError(
                "no extractable text — this looks like a scanned PDF, which needs OCR"
            )
        return text

    if name.endswith(".docx"):
        try:
            import docx
        except ImportError:
            raise ExtractError("python-docx is not installed")
        try:
            d = docx.Document(io.BytesIO(data))
        except Exception as exc:
            raise ExtractError(f"could not read DOCX: {exc}")
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    if name.endswith(".doc"):
        raise ExtractError("legacy .doc is not supported — save as .docx or PDF")

    raise ExtractError(f"unsupported file type: {filename}")


def split_sections(text: str) -> list[dict]:
    """Split into [{heading, level, content}], preserving document order."""
    lines = text.splitlines()
    marks: list[tuple[int, str, int]] = []  # (line index, heading, level)

    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            continue

        md = _MD_HEADING.match(line)
        if md:
            marks.append((i, md.group(2).strip(), len(md.group(1))))
            continue

        num = _NUM_HEADING.match(stripped)
        if num:
            level = num.group(1).count(".") + 1
            marks.append((i, f"{num.group(1)} {num.group(2)}".strip(), level))

    if not marks:
        return _fallback_chunks(text)

    sections = []
    preamble = "\n".join(lines[: marks[0][0]]).strip()
    if preamble:
        sections.append({"heading": "(document start)", "level": 1, "content": preamble})

    for idx, (line_no, heading, level) in enumerate(marks):
        end = marks[idx + 1][0] if idx + 1 < len(marks) else len(lines)
        body = "\n".join(lines[line_no:end]).strip()
        if body:
            sections.append({"heading": heading, "level": level, "content": body})

    # A single giant section is useless for attachment; break it down further.
    out: list[dict] = []
    for s in sections:
        if len(s["content"]) > FALLBACK_CHUNK_CHARS * 3:
            for j, piece in enumerate(_split_chars(s["content"])):
                out.append({
                    "heading": f"{s['heading']} (part {j + 1})",
                    "level": s["level"],
                    "content": piece,
                })
        else:
            out.append(s)
    return out


def _split_chars(text: str) -> list[str]:
    """Chunk on paragraph boundaries, never mid-paragraph."""
    paras = text.split("\n\n")
    chunks, cur = [], ""
    for p in paras:
        if cur and len(cur) + len(p) + 2 > FALLBACK_CHUNK_CHARS:
            chunks.append(cur.strip())
            cur = p
        else:
            cur = f"{cur}\n\n{p}" if cur else p
    if cur.strip():
        chunks.append(cur.strip())
    return chunks or [text]


def _fallback_chunks(text: str) -> list[dict]:
    return [
        {"heading": f"Part {i + 1}", "level": 1, "content": c}
        for i, c in enumerate(_split_chars(text))
    ]
